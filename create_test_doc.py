from docx import Document
import os

try:
    doc = Document()
    doc.add_paragraph("This is page 1.")
    doc.add_page_break()
    doc.add_paragraph("This is page 2.")
    doc.add_page_break()
    doc.add_paragraph("This is page 3.")
    
    doc.save("simple_test.docx")
    print("Created simple_test.docx")
except Exception as e:
    print(f"Error creating doc: {e}")
